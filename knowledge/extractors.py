"""多格式文档文本提取：PDF / Word / Excel / Markdown / 纯文本。

统一入口是 extract_text(path)，按扩展名分发到对应解析器。
以后要支持新格式（如 epub、html），只需写一个提取函数并注册进 EXTRACTORS。

各格式的取舍说明：
- .pdf  用 pypdf 提取文字层；扫描件（图片型 PDF）没有文字层，需要 OCR，暂不支持；
- .docx 用 python-docx，正文段落 + 表格都会提取；
- .doc  是老格式，只能借助本机安装的 Word（通过 pywin32 调用）读取；
- .xlsx 用 openpyxl，.xls 用 xlrd；每个工作表按行拼成 " | " 分隔的文本。
"""
from pathlib import Path

MAX_SHEET_ROWS = 2000  # 单个工作表最多提取的行数，防止超大表格撑爆内存


def _extract_txt(path):
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(path):
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"[第{i}页]\n{text}")
    if not parts:
        raise ValueError(
            "PDF 未提取到文字，可能是扫描件（图片型 PDF），需要 OCR，本项目暂不支持"
        )
    return "\n\n".join(parts)


def _extract_docx(path):
    import docx

    doc = docx.Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:  # 表格也提取，按行拼接
        for row in table.rows:
            line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if line:
                parts.append(line)
    if not parts:
        raise ValueError("DOCX 中未提取到文字")
    return "\n\n".join(parts)


def _extract_doc(path):
    try:
        import win32com.client
    except ImportError:
        raise ValueError(
            "读取 .doc 需要安装 pywin32 且本机装有 Word；"
            "更简单的做法是用 Word 把它另存为 .docx 后放入"
        )
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
    finally:
        word.Quit()
    if not text.strip():
        raise ValueError("DOC 中未提取到文字")
    return text


def _rows_to_text(sheet_name, rows):
    """把表格行列表拼成带工作表名的文本块。"""
    lines = []
    for row in rows[:MAX_SHEET_ROWS]:
        line = " | ".join(str(c).strip() for c in row if c is not None and str(c).strip())
        if line:
            lines.append(line)
    if not lines:
        return ""
    return f"[工作表：{sheet_name}]\n" + "\n".join(lines)


def _extract_xlsx(path):
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        parts = [_rows_to_text(ws.title, list(ws.iter_rows(values_only=True))[: MAX_SHEET_ROWS + 1]) for ws in wb.worksheets]
    finally:
        wb.close()
    result = "\n\n".join(p for p in parts if p)
    if not result:
        raise ValueError("XLSX 中未提取到内容")
    return result


def _extract_xls(path):
    import xlrd

    wb = xlrd.open_workbook(str(path))
    parts = []
    for ws in wb.sheets():
        rows = [[c.value for c in ws.row(r)] for r in range(ws.nrows)]
        parts.append(_rows_to_text(ws.name, rows))
    result = "\n\n".join(p for p in parts if p)
    if not result:
        raise ValueError("XLS 中未提取到内容")
    return result


EXTRACTORS = {
    ".md": _extract_txt,
    ".txt": _extract_txt,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".doc": _extract_doc,
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xls,
}

SUPPORTED_EXTS = tuple(EXTRACTORS)


def extract_text(path):
    """提取任意受支持文档的全部文本。path 为 pathlib.Path。"""
    ext = path.suffix.lower()
    fn = EXTRACTORS.get(ext)
    if fn is None:
        raise ValueError(f"不支持的格式：{ext}（支持：{'/'.join(SUPPORTED_EXTS)}）")
    return fn(path)
